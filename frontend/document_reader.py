"""本地学习资料文本提取。"""
from html.parser import HTMLParser
from pathlib import Path
import re


SUPPORTED_EXTENSIONS = {
    '.txt', '.md', '.markdown', '.rst', '.csv', '.json', '.py',
    '.html', '.htm', '.pdf', '.docx',
}
MAX_FILE_SIZE = 20 * 1024 * 1024


class DocumentReadError(ValueError):
    """资料无法读取或没有可用文本。"""


class _HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []

    def handle_data(self, data):
        if data.strip():
            self.parts.append(data.strip())


def _read_text_file(path):
    raw = path.read_bytes()
    for encoding in ('utf-8-sig', 'utf-8', 'gb18030'):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentReadError('无法识别文本编码，请将文件转换为 UTF-8')


def _read_pdf(path):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentReadError('缺少 pypdf，请重新安装 requirements.txt') from exc
    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                reader.decrypt('')
            except Exception as exc:
                raise DocumentReadError('PDF 已加密，无法读取') from exc
        return '\n\n'.join((page.extract_text() or '') for page in reader.pages)
    except DocumentReadError:
        raise
    except Exception as exc:
        raise DocumentReadError(f'PDF 解析失败：{exc}') from exc


def _read_docx(path):
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentReadError('缺少 python-docx，请重新安装 requirements.txt') from exc
    try:
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append('\t'.join(cell.text for cell in row.cells))
        return '\n'.join(parts)
    except Exception as exc:
        raise DocumentReadError(f'DOCX 解析失败：{exc}') from exc


def _normalize_text(text):
    text = text.replace('\x00', '').replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def read_document(file_path):
    """读取资料，返回文件信息和规范化后的正文。"""
    path = Path(file_path)
    if not path.is_file():
        raise DocumentReadError('文件不存在')
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentReadError(f'暂不支持 {suffix or "无扩展名"} 文件')
    if path.stat().st_size > MAX_FILE_SIZE:
        raise DocumentReadError('文件超过 20MB，请拆分后再导入')

    if suffix == '.pdf':
        text = _read_pdf(path)
    elif suffix == '.docx':
        text = _read_docx(path)
    else:
        text = _read_text_file(path)
        if suffix in ('.html', '.htm'):
            parser = _HTMLTextExtractor()
            parser.feed(text)
            text = '\n'.join(parser.parts)

    text = _normalize_text(text)
    if len(text) < 30:
        hint = '；扫描版 PDF 需要先进行 OCR' if suffix == '.pdf' else ''
        raise DocumentReadError(f'未提取到足够的文字内容{hint}')
    return {
        'filename': path.name,
        'file_type': suffix.lstrip('.'),
        'content': text,
        'char_count': len(text),
    }
